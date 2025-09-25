"""
Document parser for extracting text from various file formats.

This module provides functionality to parse and extract text content from
PDF, DOCX, TXT, and Markdown files for the RAG English Study system.
"""

import logging
from pathlib import Path
from typing import Optional, Dict, Any
import mimetypes
import hashlib
from datetime import datetime

# Document parsing imports (will be available after dependency installation)
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

from ..models.document import Document


logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """문서 파싱 중 발생하는 오류를 나타내는 예외 클래스."""
    pass


class DocumentParser:
    """다양한 파일 형식의 문서를 파싱하여 텍스트를 추출하는 클래스.
    
    지원하는 파일 형식:
    - PDF (.pdf)
    - Microsoft Word (.docx)
    - 텍스트 파일 (.txt)
    - Markdown 파일 (.md)
    """
    
    # 지원하는 파일 확장자와 MIME 타입 매핑
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.txt': 'text/plain',
        '.md': 'text/markdown'
    }
    
    def __init__(self):
        """DocumentParser 인스턴스를 초기화합니다."""
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """필요한 의존성 패키지들이 설치되어 있는지 확인합니다."""
        missing_deps = []
        
        if not PDF_AVAILABLE:
            missing_deps.append("pypdf (for PDF parsing)")
        if not DOCX_AVAILABLE:
            missing_deps.append("python-docx (for DOCX parsing)")
        if not MARKDOWN_AVAILABLE:
            missing_deps.append("markdown (for Markdown parsing)")
        
        if missing_deps:
            logger.warning(f"Missing dependencies: {', '.join(missing_deps)}")
    
    def is_supported_file(self, file_path: str) -> bool:
        """파일이 지원되는 형식인지 확인합니다.
        
        Args:
            file_path: 확인할 파일의 경로
            
        Returns:
            지원되는 파일 형식이면 True, 아니면 False
        """
        try:
            path = Path(file_path)
            return path.suffix.lower() in self.SUPPORTED_EXTENSIONS
        except Exception as e:
            logger.error(f"Error checking file support for {file_path}: {e}")
            return False
    
    def get_file_type(self, file_path: str) -> Optional[str]:
        """파일의 타입을 확인합니다.
        
        Args:
            file_path: 파일 경로
            
        Returns:
            파일 타입 (pdf, docx, txt, md) 또는 None
        """
        try:
            path = Path(file_path)
            extension = path.suffix.lower()
            
            if extension == '.pdf':
                return 'pdf'
            elif extension == '.docx':
                return 'docx'
            elif extension == '.txt':
                return 'txt'
            elif extension == '.md':
                return 'md'
            else:
                return None
        except Exception as e:
            logger.error(f"Error determining file type for {file_path}: {e}")
            return None
    
    def validate_file(self, file_path: str) -> bool:
        """파일이 존재하고 읽을 수 있는지 검증합니다.
        
        Args:
            file_path: 검증할 파일 경로
            
        Returns:
            파일이 유효하면 True, 아니면 False
        """
        try:
            path = Path(file_path)
            
            # 파일 존재 여부 확인
            if not path.exists():
                logger.error(f"File does not exist: {file_path}")
                return False
            
            # 파일인지 확인 (디렉토리가 아닌)
            if not path.is_file():
                logger.error(f"Path is not a file: {file_path}")
                return False
            
            # 읽기 권한 확인
            if not path.stat().st_size > 0:
                logger.warning(f"File is empty: {file_path}")
                return False
            
            # 지원되는 파일 형식인지 확인
            if not self.is_supported_file(file_path):
                logger.error(f"Unsupported file type: {file_path}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error validating file {file_path}: {e}")
            return False
    
    def parse_file(self, file_path: str) -> Optional[Document]:
        """파일을 파싱하여 Document 객체를 생성합니다.
        
        Args:
            file_path: 파싱할 파일의 경로
            
        Returns:
            파싱된 Document 객체 또는 None (실패 시)
            
        Raises:
            DocumentParsingError: 파싱 중 오류가 발생한 경우
        """
        try:
            # 파일 검증
            if not self.validate_file(file_path):
                raise DocumentParsingError(f"File validation failed: {file_path}")
            
            # 파일 타입 확인
            file_type = self.get_file_type(file_path)
            if not file_type:
                raise DocumentParsingError(f"Unsupported file type: {file_path}")
            
            # 파일 타입별 텍스트 추출
            if file_type == 'pdf':
                content = self.extract_text_from_pdf(file_path)
            elif file_type == 'docx':
                content = self.extract_text_from_docx(file_path)
            elif file_type == 'txt':
                content = self.extract_text_from_txt(file_path)
            elif file_type == 'md':
                content = self.extract_text_from_md(file_path)
            else:
                raise DocumentParsingError(f"No parser available for file type: {file_type}")
            
            if not content or not content.strip():
                raise DocumentParsingError(f"No content extracted from file: {file_path}")
            
            # Document 객체 생성
            path = Path(file_path)
            document_id = self._generate_document_id(file_path)
            title = self._extract_title(path.stem, content)
            
            document = Document(
                id=document_id,
                title=title,
                file_path=str(path.absolute()),
                content=content.strip(),
                file_type=file_type,
                created_at=datetime.now(),
                word_count=len(content.split()),
                language="english",  # 기본값, 추후 언어 감지 기능 추가 가능
                file_hash=self._generate_file_hash(file_path)
            )
            
            logger.info(f"Successfully parsed document: {file_path} ({document.word_count} words)")
            return document
            
        except Exception as e:
            error_msg = f"Failed to parse file {file_path}: {str(e)}"
            logger.error(error_msg)
            raise DocumentParsingError(error_msg) from e
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """PDF 파일에서 텍스트를 추출합니다.
        
        Args:
            file_path: PDF 파일 경로
            
        Returns:
            추출된 텍스트 내용
            
        Raises:
            DocumentParsingError: PDF 파싱 실패 시
        """
        if not PDF_AVAILABLE:
            raise DocumentParsingError("pypdf package is required for PDF parsing")
        
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # 각 페이지에서 텍스트 추출
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1} in {file_path}: {e}")
                        continue
            
            if not text_content:
                raise DocumentParsingError(f"No text content found in PDF: {file_path}")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise DocumentParsingError(f"PDF parsing failed for {file_path}: {str(e)}") from e
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """DOCX 파일에서 텍스트를 추출합니다.
        
        Args:
            file_path: DOCX 파일 경로
            
        Returns:
            추출된 텍스트 내용
            
        Raises:
            DocumentParsingError: DOCX 파싱 실패 시
        """
        if not DOCX_AVAILABLE:
            raise DocumentParsingError("python-docx package is required for DOCX parsing")
        
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # 각 문단에서 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 표에서 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise DocumentParsingError(f"No text content found in DOCX: {file_path}")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise DocumentParsingError(f"DOCX parsing failed for {file_path}: {str(e)}") from e
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """텍스트 파일에서 내용을 읽습니다.
        
        Args:
            file_path: 텍스트 파일 경로
            
        Returns:
            파일 내용
            
        Raises:
            DocumentParsingError: 텍스트 파일 읽기 실패 시
        """
        try:
            # 다양한 인코딩으로 시도
            encodings = ['utf-8', 'utf-8-sig', 'cp949', 'euc-kr', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content.strip():
                            return content
                except UnicodeDecodeError:
                    continue
            
            raise DocumentParsingError(f"Could not decode text file with any supported encoding: {file_path}")
            
        except Exception as e:
            raise DocumentParsingError(f"Text file parsing failed for {file_path}: {str(e)}") from e
    
    def extract_text_from_md(self, file_path: str) -> str:
        """Markdown 파일에서 텍스트를 추출합니다.
        
        Args:
            file_path: Markdown 파일 경로
            
        Returns:
            추출된 텍스트 내용 (HTML 태그 제거됨)
            
        Raises:
            DocumentParsingError: Markdown 파싱 실패 시
        """
        try:
            # 먼저 원본 텍스트를 읽음
            raw_content = self.extract_text_from_txt(file_path)
            
            if MARKDOWN_AVAILABLE:
                # Markdown을 HTML로 변환한 후 텍스트만 추출
                md = markdown.Markdown()
                html_content = md.convert(raw_content)
                
                # 간단한 HTML 태그 제거 (더 정교한 처리를 위해서는 BeautifulSoup 사용 가능)
                import re
                text_content = re.sub(r'<[^>]+>', '', html_content)
                text_content = re.sub(r'\s+', ' ', text_content).strip()
                
                return text_content if text_content else raw_content
            else:
                # Markdown 패키지가 없으면 원본 텍스트 반환
                logger.warning(f"Markdown package not available, returning raw content for {file_path}")
                return raw_content
            
        except Exception as e:
            raise DocumentParsingError(f"Markdown parsing failed for {file_path}: {str(e)}") from e
    
    def _generate_document_id(self, file_path: str) -> str:
        """파일 경로를 기반으로 고유한 문서 ID를 생성합니다.
        
        Args:
            file_path: 파일 경로
            
        Returns:
            생성된 문서 ID
        """
        # 파일 경로의 해시값을 사용하여 고유 ID 생성
        path_hash = hashlib.md5(str(Path(file_path).absolute()).encode('utf-8')).hexdigest()
        return f"doc_{path_hash[:12]}"
    
    def _generate_file_hash(self, file_path: str) -> str:
        """파일 내용의 해시값을 생성합니다.
        
        Args:
            file_path: 파일 경로
            
        Returns:
            파일 내용의 MD5 해시값
        """
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.error(f"Failed to generate file hash for {file_path}: {e}")
            return ""
    
    def _extract_title(self, filename: str, content: str) -> str:
        """파일명과 내용을 기반으로 문서 제목을 추출합니다.
        
        Args:
            filename: 파일명 (확장자 제외)
            content: 문서 내용
            
        Returns:
            추출된 제목
        """
        # 첫 번째 줄이 제목처럼 보이면 사용
        lines = content.strip().split('\n')
        if lines:
            first_line = lines[0].strip()
            # 첫 줄이 너무 길지 않고 의미있는 내용이면 제목으로 사용
            if len(first_line) <= 100 and len(first_line.split()) <= 15:
                # Markdown 헤더 표시 제거
                title = first_line.lstrip('#').strip()
                if title:
                    return title
        
        # 파일명을 제목으로 사용 (언더스코어를 공백으로 변경)
        return filename.replace('_', ' ').replace('-', ' ').title()
    
    def get_supported_extensions(self) -> list[str]:
        """지원되는 파일 확장자 목록을 반환합니다.
        
        Returns:
            지원되는 파일 확장자 리스트
        """
        return list(self.SUPPORTED_EXTENSIONS.keys())
    
    def get_parser_info(self) -> Dict[str, Any]:
        """파서의 현재 상태와 지원 정보를 반환합니다.
        
        Returns:
            파서 정보 딕셔너리
        """
        return {
            'supported_extensions': self.get_supported_extensions(),
            'dependencies': {
                'pypdf': PDF_AVAILABLE,
                'python-docx': DOCX_AVAILABLE,
                'markdown': MARKDOWN_AVAILABLE
            },
            'parser_version': '1.0.0'
        }